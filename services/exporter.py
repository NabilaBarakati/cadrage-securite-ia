import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, HRFlowable
from reportlab.lib import colors


def export_to_docx(md_content: str) -> bytes:
    doc = Document()

    doc.core_properties.author = "Cadrage Sécurité IA"
    doc.core_properties.created = datetime.now()

    for line in md_content.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.strip() == "---":
            doc.add_paragraph("─" * 50)
        elif line.strip():
            # Strip inline bold/italic markers
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            doc.add_paragraph(clean)
        else:
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_to_pdf(md_content: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    style_h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    style_h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, spaceAfter=8)
    style_h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=12, spaceAfter=6)
    style_body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, spaceAfter=4, leading=14)

    story = []

    for line in md_content.splitlines():
        if line.startswith("# "):
            story.append(Paragraph(line[2:], style_h1))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], style_h2))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], style_h3))
        elif line.strip() == "---":
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", color=colors.grey))
            story.append(Spacer(1, 4))
        elif line.strip():
            clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
            clean = re.sub(r"\*(.+?)\*", r"<i>\1</i>", clean)
            story.append(Paragraph(clean, style_body))
        else:
            story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()
